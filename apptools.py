from basictools import *
import shutil
import os
from constants import APP_ALIASES, BASE_DIR, NOTES_DIR, SEARCH_DIRS

def excel_create(filename, headers=None, subfolder=""):
    """Create a new empty Excel file, optionally with header row.

    Args:
        filename: Just the filename e.g. data.xlsx
        headers: Optional list of column headers e.g. ["Name", "Age", "Email"]
        subfolder: Optional subfolder inside agent folder

    Returns:
        Confirmation message with full path
    """
    try:
        import openpyxl
    except ImportError:
        return "Error: openpyxl is not installed. Run: pip install openpyxl"

    base = os.path.join(BASE_DIR, subfolder) if subfolder else NOTES_DIR
    os.makedirs(base, exist_ok=True)
    full_path = os.path.join(base, os.path.basename(filename))

    wb = openpyxl.Workbook()
    ws = wb.active
    if headers:
        ws.append(headers)
    wb.save(full_path)
    return f"Excel file created: {full_path}"

def excel_read(filename, subfolder="", sheet=None):
    """Read all cell values from an Excel file (.xlsx).

    Args:
        filename: Just the filename e.g. data.xlsx
        subfolder: Optional subfolder hint
        sheet: Sheet name to read — leave empty for the first sheet

    Returns:
        A text table of the sheet contents
    """
    try:
        import openpyxl
    except ImportError:
        return "Error: openpyxl is not installed. Run: pip install openpyxl"

    full_path = os.path.join(BASE_DIR, subfolder, filename) if subfolder else find_file(filename)
    if not full_path or not os.path.exists(full_path):
        return f"Error: {filename} not found"

    wb = openpyxl.load_workbook(full_path, data_only=True)
    ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active

    rows = []
    for row in ws.iter_rows(values_only=True):
        rows.append("\t".join("" if v is None else str(v) for v in row))
    return f"Sheet: {ws.title}\n" + "\n".join(rows) if rows else "Empty sheet"


def excel_write_cell(filename, cell, value, subfolder="", sheet=None):
    """Write a value to a specific cell in an Excel file.
    Creates the file if it doesn't exist.

    Args:
        filename: Just the filename e.g. data.xlsx
        cell: Cell reference e.g. A1, B3
        value: Value to write into the cell
        subfolder: Optional subfolder
        sheet: Sheet name — leave empty for the first sheet

    Returns:
        Confirmation message
    """
    try:
        import openpyxl
    except ImportError:
        return "Error: openpyxl is not installed. Run: pip install openpyxl"

    base = os.path.join(BASE_DIR, subfolder) if subfolder else NOTES_DIR
    os.makedirs(base, exist_ok=True)
    full_path = os.path.join(base, os.path.basename(filename))

    if os.path.exists(full_path):
        wb = openpyxl.load_workbook(full_path)
    else:
        wb = openpyxl.Workbook()

    ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active
    ws[cell] = value
    wb.save(full_path)
    return f"Written '{value}' to {cell} in {full_path}"


def excel_add_row(filename, row_data, subfolder="", sheet=None):
    """Append a new row of data to an Excel file.
    Creates the file if it doesn't exist.

    Args:
        filename: Just the filename e.g. data.xlsx
        row_data: List of values for the new row e.g. ["Alice", 30, "Engineer"]
        subfolder: Optional subfolder
        sheet: Sheet name — leave empty for the first sheet

    Returns:
        Confirmation message
    """
    try:
        import openpyxl
    except ImportError:
        return "Error: openpyxl is not installed. Run: pip install openpyxl"

    base = os.path.join(BASE_DIR, subfolder) if subfolder else NOTES_DIR
    os.makedirs(base, exist_ok=True)
    full_path = os.path.join(base, os.path.basename(filename))

    if os.path.exists(full_path):
        wb = openpyxl.load_workbook(full_path)
    else:
        wb = openpyxl.Workbook()

    ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active
    ws.append(row_data)
    wb.save(full_path)
    return f"Row {row_data} appended to {full_path}"

def word_create(filename, content="", subfolder=""):
    """Create a new Word document (.docx) with optional initial content.

    Args:
        filename: Just the filename e.g. report.docx
        content: Optional initial paragraph text
        subfolder: Optional subfolder

    Returns:
        Confirmation message with full path
    """
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx is not installed. Run: pip install python-docx"

    base = os.path.join(BASE_DIR, subfolder) if subfolder else NOTES_DIR
    os.makedirs(base, exist_ok=True)
    full_path = os.path.join(base, os.path.basename(filename))

    doc = Document()
    if content:
        doc.add_paragraph(content)
    doc.save(full_path)
    return f"Word document created: {full_path}"


def word_read(filename, subfolder=""):
    """Read all text from a Word document (.docx).

    Args:
        filename: Just the filename e.g. report.docx
        subfolder: Optional subfolder hint

    Returns:
        The full text content of the document
    """
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx is not installed. Run: pip install python-docx"

    full_path = os.path.join(BASE_DIR, subfolder, filename) if subfolder else find_file(filename)
    if not full_path or not os.path.exists(full_path):
        return f"Error: {filename} not found"

    doc = Document(full_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    return text if text.strip() else "Document is empty"


def word_append(filename, text, subfolder=""):
    """Append a new paragraph to an existing Word document.

    Args:
        filename: Just the filename e.g. report.docx
        text: The paragraph text to append
        subfolder: Optional subfolder hint

    Returns:
        Confirmation message
    """
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx is not installed. Run: pip install python-docx"

    full_path = os.path.join(BASE_DIR, subfolder, filename) if subfolder else find_file(filename)
    if not full_path or not os.path.exists(full_path):
        return f"Error: {filename} not found"

    doc = Document(full_path)
    doc.add_paragraph(text)
    doc.save(full_path)
    return f"Paragraph appended to {full_path}"


def word_replace(filename, find, replace, subfolder=""):
    """Find and replace text throughout a Word document.

    Args:
        filename: Just the filename e.g. report.docx
        find: The text to search for
        replace: The text to replace it with
        subfolder: Optional subfolder hint

    Returns:
        Confirmation message with count of replacements made
    """
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx is not installed. Run: pip install python-docx"
 
    full_path = os.path.join(BASE_DIR, subfolder, filename) if subfolder else find_file(filename)
    if not full_path or not os.path.exists(full_path):
        return f"Error: {filename} not found"
 
    doc = Document(full_path)
    count = 0
    for para in doc.paragraphs:
        if find in para.text:
            for run in para.runs:
                if find in run.text:
                    run.text = run.text.replace(find, replace)
                    count += 1
    doc.save(full_path)
    return f"Replaced {count} occurrence(s) of '{find}' with '{replace}' in {full_path}"